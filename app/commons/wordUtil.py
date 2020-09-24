import os
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION, WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import RGBColor
from docx.shared import Pt
from docx.oxml.ns import qn

from pylab import mpl
import numpy as np

# 解决中文乱码问题
mpl.rcParams['font.sans-serif'] = ['FangSong']  # 指定默认字体
mpl.rcParams['axes.unicode_minus'] = False  # 解决保存图像是负号'-'显示为方块的问题


class WordUtils:
    def __init__(self, fileName):
        self._file = fileName
        self._doc = Document()

        # 设置默认字体
        self._doc.styles['Normal'].font.name = '微软雅黑'
        self._doc.styles['Normal']._element.rPr.rFonts.set(
            qn('w:eastAsia'), '微软雅黑')

    # 添加段落
    def _add_paragraph(self,
                       content,
                       style="ListBullet",
                       bold=False,
                       indent=False,
                       right=False,
                       size=None):
        paragraph = self._doc.add_paragraph()
        run = paragraph.add_run(content)
        run.bold = bold
        if size:
            run.font.size = Inches(size / 40)
        if right:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        else:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if indent:
            paragraph.paragraph_format.left_indent = Inches(1.25)
        return paragraph

    # 添加标题
    def _add_heading(self, content, level=2, size=12):
        head = self._doc.add_heading(level=level)
        run = head.add_run(content)
        run.font.size = Pt(24)
        run.font.color.rgb = RGBColor(0, 0, 0)
        head.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 添加分页符
    def _add_page_break(self):
        self._doc.add_page_break()

    # 添加表格
    def _add_table(self, rows, cols, data=[[]], style="Table Grid"):
        table = self._doc.add_table(rows, cols)
        table.style = style
        table.autofit = True
        rows = len(data)
        for i in range(rows):
            for j in range(cols):
                table.cell(i, j).text = str(data[i][j])
                table.cell(i, j).paragraphs[
                    0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        return table

    def _merge(self, table, row1, col1, row2, col2):
        table.cell(row1, col1).merge(table.cell(row2, col2))

    # 添加图片
    def _add_picture(self, location, width=1):
        self._doc.add_picture(location, width=Inches(width))

    # 添加章节
    def _add_section(self):
        return self._doc.add_section(WD_SECTION.ODD_PAGE)

    def get_picture(self, filename, xdata=None, ydata=None, title=None):
        import matplotlib.pyplot as plt

        fig, ax = plt.subplots()
        index = np.arange(len(ydata))
        ax.set_title(title)
        ax.set_xticks(ticks=index)
        rect1 = ax.bar(
            range(len(ydata)),
            ydata,
            0.2,
            tick_label=xdata)

        def auto_text(rects):
            for rect in rects:
                ax.text(rect.get_x(), rect.get_height(), rect.get_height(), ha='left', va='bottom')

        plt.legend()
        auto_text(rect1)
        plt.savefig(filename)

    # 保存文件
    def _save(self):
        self._doc.save(self._file)


def test():
    a = WordUtils("test.docx")
    # a._add_heading('Test Heading')
    # section = a._add_section()
    # section.orientation = WD_ORIENT.LANDSCAPE
    # # section.page_width = 10
    # # section.page_height = 75
    # a._add_paragraph("test paragraph")
    # a._add_page_break()
    # tb1 = a._add_table(3, 2, [["name", "age"], ["", "12"], ["das2", "43"]])
    # a._merge(tb1, 0, 0, 1, 0)
    # a._save()
    a.get_picture(["yy", "xx", "yy"], [111, 122, 324])


if __name__ == "__main__":
    test()
